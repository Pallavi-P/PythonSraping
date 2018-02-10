from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import requests
import os
import sys
import re


index = 0
limit = int(input("Maximum number of url's you want to display:"))

r  = requests.get("https://www.reliablesoft.net/top-10-search-engines-in-the-world/")
data = r.text


soup = BeautifulSoup(data, "lxml")
document = Document()
document.add_heading('List of urls')

yes = {'yes','y'}
no = {'no','n'}

choice = input('Would you like to download Docx file Please respond with yes or no:').lower()
if choice in yes:
        try:
                for link in soup.find_all('a'):
                        pattern = (r'http[s]?://(?:[a-z]*\.(google|bing)\.)+')
                        if re.match(pattern, link.get('href')) is not None:

                                # Adding data to docx file
                                document.add_paragraph(link.get('href'))
                                index += 1
                                if index == limit:
                                        break
                document.save('scrap.docx')
                print('Please check the docx file with name scrape.docx')

        except OSError:
                pass

elif choice in no:
                for link in soup.find_all('a'):
                        #print(link.get('href'))
                        pattern = (r'http[s]?://(?:[a-z]*\.(google|bing)\.)+')
                        if re.match(pattern, link.get('href')) is not None:
                                print(link.get('href'))
                                index += 1
                                if index == limit:
                                        break
else:
   sys.stdout.write("Please respond with 'yes' or 'no' after running the program again")
